
import streamlit as st
import docx
from io import BytesIO

# Dicionário com competências e habilidades (resumido para exemplo)
competencias = {
    "CE1": {
        "descricao": "Analisar processos políticos, econômicos, sociais, ambientais e culturais nos âmbitos local, regional, nacional e mundial.",
        "habilidades": {
            "EM13CHS101": "Identificar, analisar e comparar diferentes fontes e narrativas expressas em diversas linguagens.",
            "EM13CHS102": "Analisar circunstâncias históricas de matrizes conceituais como etnocentrismo, racismo, etc.",
            "EM13CHS103": "Elaborar hipóteses e compor argumentos com base em dados e evidências.",
            "EM13CHS104": "Analisar objetos e vestígios da cultura material e imaterial.",
            "EM13CHS105": "Criticar oposições dicotômicas como cidade/campo, civilizados/bárbaros, etc.",
            "EM13CHS106": "Utilizar linguagens cartográfica e digital de forma crítica e ética."
        }
    },
    "CE2": {
        "descricao": "Analisar a formação de territórios e fronteiras, e o papel geopolítico dos Estados-nações.",
        "habilidades": {
            "EM13CHS201": "Analisar dinâmicas populacionais e mobilidades humanas.",
            "EM13CHS202": "Analisar impactos das tecnologias nas sociedades.",
            "EM13CHS203": "Comparar significados de território, fronteira e vazio.",
            "EM13CHS204": "Avaliar processos de ocupação e conflitos territoriais.",
            "EM13CHS205": "Analisar a produção de territorialidades contemporâneas.",
            "EM13CHS206": "Aplicar princípios geográficos na análise da ocupação do espaço."
        }
    },
    # Demais competências podem ser adicionadas aqui da mesma forma
}

st.set_page_config(page_title="BNCC: Competências e Habilidades", layout="centered")
st.title("🎓 Seletor de Competências e Habilidades - BNCC")
st.markdown("Escolha as competências e as habilidades que deseja incluir no seu documento.")

escolhas = {}

# Interface de seleção
for cod, comp in competencias.items():
    with st.expander(f"{cod} - {comp['descricao']}"):
        habilidades_escolhidas = st.multiselect(
            f"Selecione as habilidades para {cod}:",
            options=list(comp["habilidades"].keys()),
            format_func=lambda x: f"{x} - {comp['habilidades'][x]}"
        )
        if habilidades_escolhidas:
            escolhas[cod] = habilidades_escolhidas

# Função para gerar o arquivo Word
def gerar_documento(selecionadas):
    doc = docx.Document()
    doc.add_heading("Competências e Habilidades Selecionadas - BNCC", 0)

    for cod, habilidades in selecionadas.items():
        doc.add_heading(f"{cod} - {competencias[cod]['descricao']}", level=1)
        for hab in habilidades:
            desc = competencias[cod]['habilidades'][hab]
            doc.add_paragraph(f"{hab}: {desc}", style='List Bullet')

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Botão para gerar e baixar o documento
if st.button("📄 Gerar Documento Word"):
    if escolhas:
        doc_file = gerar_documento(escolhas)
        st.download_button(
            label="📥 Baixar Documento",
            data=doc_file,
            file_name="competencias_habilidades_bncc.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    else:
        st.warning("Por favor, selecione pelo menos uma habilidade.")
